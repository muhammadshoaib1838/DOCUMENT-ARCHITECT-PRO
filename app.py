import easyocr
import cv2
import numpy as np
import gradio as gr
from groq import Groq
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# --- API Key handling ---
# Prioritizes environment variables (ideal for HF Spaces/Docker)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")

# --- Engines ---
# Note: EasyOCR is set to CPU mode (gpu=False) for better compatibility
reader = easyocr.Reader(['en'], gpu=False)
client = Groq(api_key=GROQ_API_KEY)

# --- Main Pipeline Logic ---
def process_pipeline(image_path, history):
    if not image_path:
        return "⚠️ Please upload an image.", "Waiting for input...", "No data", None, None, history

    if not GROQ_API_KEY:
        return "❌ API Key Missing", "Please set GROQ_API_KEY in environment variables.", "N/A", None, None, history

    try:
        # 1. Image Preprocessing & OCR
        img = cv2.imread(image_path)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        processed_img = cv2.adaptiveThreshold(
            cv2.GaussianBlur(gray, (5, 5), 0),
            255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        raw_text = " ".join(reader.readtext(processed_img, detail=0))

        if not raw_text.strip():
            return "❌ No text detected.", "Try a clearer image.", "N/A", None, None, history

        # 2. AI Structuring (Groq)
        prompt = f"""
        Act as a Professional Document Architect.
        Convert this OCR text into a beautiful digital document:
        - Use # for the Main Title
        - Use ## for Section Headings
        - Use bolding (**) for important technical terms
        - End with a section titled '--- FINAL SUMMARY ---'

        TEXT: {raw_text}
        """

        chat = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile"
        )
        full_output = chat.choices[0].message.content

        if "--- FINAL SUMMARY ---" in full_output:
            notes, summary = full_output.split("--- FINAL SUMMARY ---")
        else:
            notes, summary = full_output, "Summary included in text."

        # 3. History Management
        new_entry = f"Note {len(history) + 1}: {notes[:30].strip()}..."
        history.append(new_entry)

        # 4. Export Files
        doc_path, pdf_path = "Architect_Export.docx", "Architect_Export.pdf"
        d = Document()
        d.add_heading('Document Architect Pro Export', 0)
        d.add_paragraph(notes)
        d.save(doc_path)

        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Document Architect Pro Export")
        c.setFont("Helvetica", 10)
        y = 720
        for line in notes.split('\n')[:40]:
            c.drawString(50, y, line[:90])
            y -= 15
        c.save()

        return raw_text, notes.strip(), summary.strip(), doc_path, pdf_path, history

    except Exception as e:
        return f"Error: {str(e)}", "An error occurred.", "N/A", None, None, history

# --- High-Contrast Professional CSS ---
css = """
.gradio-container-5-50-0-dev0 .prose * { color: white !important; }
footer {visibility: hidden}
.gradio-container { background-color: #020617 !important; color: #f8fafc !important; }
.header-container { text-align: center; padding: 30px 10px; border-bottom: 1px solid #1e293b; margin-bottom: 25px; }
.header-container h1 { color: #ffffff !important; text-shadow: 0 0 10px rgba(99, 102, 241, 0.5); margin-bottom: 20px; }
.jaman-tagline {
    background-color: #4c0519 !important;
    color: #ffff00 !important;
    padding: 15px 30px;
    border-radius: 12px;
    display: inline-block;
    font-weight: 900 !important;
    font-size: 1.3em;
    letter-spacing: 0.5px;
    box-shadow: 0 4px 20px rgba(0,0,0,0.5);
    border: 1px solid rgba(255,255,255,0.2);
}
.pro-output, .pro-output *, .history-box, textarea { background-color: #0f172a !important; color: #ffffff !important; }
.pro-output h1 { color: #38bdf8 !important; border-bottom: 1px solid #1e40af; }
.pro-output h2 { color: #c084fc !important; }
.pro-output strong { color: #fbbf24 !important; }
.history-box { border: 1px solid #334155 !important; border-radius: 12px; padding: 10px; }
.history-box * { color: #ffffff !important; }
.pro-btn { background: linear-gradient(135deg, #4f46e5 0%, #9333ea 100%) !important; color: white !important; font-weight: bold !important; border: none !important; }
.signature-name { color: #c084fc !important; font-size: 1.6em !important; font-family: 'serif'; font-weight: bold; text-shadow: 0 0 8px rgba(168, 85, 247, 0.4); }
"""

with gr.Blocks(css=css) as demo:
    history_state = gr.State([])

    with gr.Column(elem_classes="header-container"):
        gr.Markdown("# 🌌 DOCUMENT ARCHITECT PRO")
        gr.Markdown("Transforming handwritten chaos into structured digital gold.", elem_classes="jaman-tagline")

    with gr.Row():
        with gr.Column(scale=1, elem_classes="history-box"):
            gr.Markdown("### 🕒 Session History")
            history_display = gr.Dataset(
                components=[gr.Textbox(visible=False)],
                label=None,
                samples=[]
            )

        with gr.Column(scale=3):
            with gr.Row():
                with gr.Column(scale=1):
                    input_img = gr.Image(type="filepath", label="📸 Upload Source")
                    submit_btn = gr.Button("ARCHITECT DOCUMENT ✨", elem_classes="pro-btn")

                with gr.Column(scale=2):
                    with gr.Tabs():
                        with gr.TabItem("✨ Digital Blueprint"):
                            output_notes = gr.Markdown(elem_classes="pro-output")
                            with gr.Row():
                                word_dl = gr.File(label="DOCX")
                                pdf_dl = gr.File(label="PDF")
                        with gr.TabItem("💡 Executive Summary"):
                            output_summary = gr.Textbox(label="TL;DR Summary", lines=8)
                        with gr.TabItem("📄 Raw OCR Buffer"):
                            output_raw = gr.Textbox(label="Initial Scan", lines=10)

    gr.HTML("""
        <div style="text-align: center; margin-top: 50px; padding: 30px; border-top: 1px solid #1e293b;">
            <p style="color: #94a3b8; font-size: 0.9em; margin-bottom: 5px; letter-spacing: 1px;">DESIGNED & ENGINEERED BY</p>
            <div class="signature-name">Muhammad Shoaib Nazz</div>
        </div>
    """)

    def run_and_update(img, hist):
        raw, notes, summary, doc, pdf, updated_hist = process_pipeline(img, hist)
        hist_samples = [[x] for x in updated_hist]
        return raw, notes, summary, doc, pdf, updated_hist, gr.Dataset(samples=hist_samples)

    submit_btn.click(
        fn=run_and_update,
        inputs=[input_img, history_state],
        outputs=[output_raw, output_notes, output_summary, word_dl, pdf_dl, history_state, history_display]
    )

if __name__ == "__main__":
    demo.launch()
